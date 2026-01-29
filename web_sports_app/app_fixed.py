from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, session
import os
import sqlite3
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import tempfile
import shutil

app = Flask(__name__)
app.secret_key = 'your_secret_key'

UPLOAD_FOLDER = 'static/uploads'
TEMPLATE_FOLDER = 'static/templates'
for folder in [UPLOAD_FOLDER, TEMPLATE_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER

DB_NAME = 'students.db'

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            dob TEXT,
            mother_name TEXT,
            father_name TEXT,
            branch_sem TEXT,
            usn TEXT UNIQUE,
            phone TEXT,
            email TEXT,
            photo_path TEXT,
            sports TEXT
        )
    ''')
    try:
        c.execute('ALTER TABLE students ADD COLUMN blood_group TEXT')
    except:
        pass
    try:
        c.execute('ALTER TABLE students ADD COLUMN gender TEXT')
    except:
        pass
    conn.commit()
    conn.close()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/index')
def index():
    return redirect(url_for('home'))

@app.route('/data-entry', methods=['GET', 'POST'])
def data_entry():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()
        blood_group = request.form.get('blood_group', '').strip()
        gender = request.form.get('gender', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))

        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                INSERT INTO students (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports, blood_group, gender)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, blood_group, gender))
            conn.commit()
            conn.close()
            flash('Student saved successfully.', 'success')
            return redirect(url_for('data_entry'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)

    return render_template('data_entry.html')

@app.route('/data-view')
def data_view():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_view.html', students=students)

@app.route('/data-edit')
def data_edit():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_edit.html', students=students)

@app.route('/edit-student/<int:student_id>', methods=['GET', 'POST'])
def edit_student(student_id):
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()
        blood_group = request.form.get('blood_group', '').strip()
        gender = request.form.get('gender', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
        else:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('SELECT photo_path FROM students WHERE id = ?', (student_id,))
            result = c.fetchone()
            conn.close()
            if result:
                photo_filename = result[0]

        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                UPDATE students SET name=?, dob=?, mother_name=?, father_name=?, branch_sem=?, usn=?, phone=?, email=?, photo_path=?, sports=?, blood_group=?, gender=?
                WHERE id=?
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, blood_group, gender, student_id))
            conn.commit()
            conn.close()
            flash('Student updated successfully.', 'success')
            return redirect(url_for('data_edit'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students WHERE id = ?', (student_id,))
    student = c.fetchone()
    conn.close()
    
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('data_edit'))
    
    return render_template('edit_student.html', student=student)

@app.route('/delete-student/<int:student_id>', methods=['POST'])
def delete_student(student_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('DELETE FROM students WHERE id = ?', (student_id,))
    conn.commit()
    conn.close()
    flash('Student deleted successfully.', 'success')
    return redirect(url_for('data_edit'))

@app.route('/data-select')
def data_select():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email, sports FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    return render_template('data_select.html', students=students)

@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/template-upload', methods=['GET', 'POST'])
def template_upload():
    if request.method == 'POST':
        template = request.files.get('template')
        if template and template.filename.endswith('.docx'):
            filename = secure_filename(template.filename)
            template_path = os.path.join(app.config['TEMPLATE_FOLDER'], filename)
            template.save(template_path)
            session['template_path'] = template_path
            session['template_uploaded'] = True
            flash('Template uploaded successfully!', 'success')
        else:
            flash('Please upload a valid .docx file.', 'error')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    
    return render_template('template_upload.html', students=students)

@app.route('/generate-report', methods=['POST'])
def generate_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    if report_format == 'vtu_eligibility':
        from docx.enum.section import WD_ORIENT
        
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('ELIGIBILITYPROFORMA')
        run.bold = True
        run.underline = True
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run('ELIGIBITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26')
        
        college_info = doc.add_paragraph()
        college_info.add_run('COLLEGE NAME & ADDRESS : ')
        run = college_info.add_run('DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082')
        run.underline = True
        
        game = doc.add_paragraph()
        game.add_run('GAME :- ____________')
        
        org_college = doc.add_paragraph()
        org_college.add_run('ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________')
        
        doc.add_paragraph('')
        
        table = doc.add_table(rows=len(students)+2, cols=7)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'A'
        hdr_cells[1].text = 'B'
        hdr_cells[2].text = 'C'
        hdr_cells[3].text = 'D'
        hdr_cells[4].text = 'E'
        hdr_cells[5].text = 'F'
        hdr_cells[6].text = 'G'
        
        hdr_cells2 = table.rows[1].cells
        hdr_cells2[0].text = 'SL NO.'
        hdr_cells2[1].text = 'Student Details'
        hdr_cells2[2].text = 'Course Details'
        hdr_cells2[3].text = 'Academic Details'
        hdr_cells2[4].text = 'VTU Previous'
        hdr_cells2[5].text = 'Photo'
        hdr_cells2[6].text = 'Signature'
        
        for i, student in enumerate(students):
            row_cells = table.rows[i+2].cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = f'Name: {student[1] or ""}\nFather: {student[4] or ""}\nMother: {student[3] or ""}\nBranch: {student[5] or ""}\nUSN: {student[6] or ""}'
            row_cells[2].text = f'Course: {student[5] or ""}\nDuration: 4 Years\nDOB: {student[2] or ""}\nContact: {student[7] or ""}'
            row_cells[3].text = 'PUC Date: ___\nFirst Admission: ___\nCurrent Admission: ___'
            row_cells[4].text = f'Game: {student[10] or ""}\nYear: ___'
            
            if student[9]:
                photo_path = os.path.join(app.config['UPLOAD_FOLDER'], student[9])
                if os.path.exists(photo_path):
                    try:
                        paragraph = row_cells[5].paragraphs[0]
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.add_picture(photo_path, width=Inches(1), height=Inches(1.3))
                    except:
                        row_cells[5].text = 'Photo Available'
                else:
                    row_cells[5].text = 'Photo Not Found'
            else:
                row_cells[5].text = 'No Photo'
            
            row_cells[6].text = ''
    
    elif report_format == 'hod_bonafide':
        for student in students:
            for i in range(10):
                doc.add_paragraph('')
            
            content = f'This is to certify that Mr/Ms {student[1] or "____________________________"} is a student of {student[5] or "___________________________"} department studying in _____________ Semester Bearing USN {student[6] or "_____________________________"} for academic year \n20__-20__.And his/her present attendance is _________% he/she can/can\'t take part in sports activity on __/__/____ to__/__/____.'
            para = doc.add_paragraph(content)
            for run in para.runs:
                run.font.size = Inches(14/72)
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            sig_para = doc.add_paragraph('Physical Education Director            Sports Co-ordinator                            Head of the Department')
            for run in sig_para.runs:
                run.font.size = Inches(14/72)
            doc.add_page_break()
    
    elif report_format == 'vtu_bonafide':
        for i in range(5):
            doc.add_paragraph('')
        
        to_para = doc.add_paragraph('To .')
        for run in to_para.runs:
            run.font.size = Inches(14/72)
        
        for i in range(3):
            doc.add_paragraph('')
        
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run('Sub : List of Students participating in ______________(game ) tournament .')
        sub_run.font.size = Inches(14/72)
        sub_run.bold = True
        
        doc.add_paragraph('')
        
        ref_para = doc.add_paragraph('With reference to the above subject , I wish to state that the following Bonafide students of our college will be participating in _____________________________________ tournament .')
        for run in ref_para.runs:
            run.font.size = Inches(14/72)
        
        doc.add_paragraph('')
        
        req_para = doc.add_paragraph('Hence, I request you to kindly permit them and oblige')
        for run in req_para.runs:
            run.font.size = Inches(14/72)
        
        doc.add_paragraph('')
        
        table = doc.add_table(rows=len(students)+1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sl.No.'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'USN'
        hdr_cells[3].text = 'Branch'
        
        for i, student in enumerate(students):
            row_cells = table.rows[i+1].cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = student[1] or ''
            row_cells[2].text = student[6] or ''
            row_cells[3].text = student[5] or ''
        
        doc.add_paragraph('')
        doc.add_paragraph('')
        
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.columns[0].width = Inches(3)
        sig_table.columns[1].width = Inches(3)
        
        left_cell = sig_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.add_run('Physical Education Director\n(for signature)')
        
        right_cell = sig_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_para.add_run('Principal\n(for signature)')
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'{report_format}_report.docx')

@app.route('/edit-report', methods=['POST'])
def edit_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    report_title = 'HOD BONAFIDE CERTIFICATE'
    report_content = '''This is to certify that Mr/Ms [NAME] is a student of [BRANCH] department studying in _____________ Semester Bearing USN [USN] for academic year 
20__-20__.And his/her present attendance is _________% he/she can/can't take part in sports activity on __/__/____ to__/__/____.



Physical Education Director            Sports Co-ordinator                            Head of the Department'''
    
    return render_template('report_editor.html', 
                         students=students, 
                         selected_students=','.join(selected_ids),
                         report_format=report_format,
                         report_title=report_title,
                         report_content=report_content)

@app.route('/generate_edited_report', methods=['POST'])
def generate_edited_report():
    selected_ids = request.form.get('selected_students').split(',')
    report_title = request.form.get('report_title')
    report_content = request.form.get('report_content')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    for student in students:
        content = report_content
        content = content.replace('[NAME]', student[1] or '')
        content = content.replace('[USN]', student[6] or '')
        content = content.replace('[BRANCH]', student[5] or '')
        
        for line in content.split('\n'):
            doc.add_paragraph(line)
        
        doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name='edited_report.docx')

@app.route('/generate-all-report', methods=['POST'])
def generate_all_report():
    report_format = request.form.get('report_format', 'detailed')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    
    if not students:
        flash('No students found to generate report.', 'error')
        return redirect(url_for('report'))
    
    doc = Document()
    
    if report_format == 'hod_bonafide':
        for student in students:
            for i in range(10):
                doc.add_paragraph('')
            
            content = f'This is to certify that Mr/Ms {student[1] or "____________________________"} is a student of {student[5] or "___________________________"} department studying in _____________ Semester Bearing USN {student[6] or "_____________________________"} for academic year \n20__-20__.And his/her present attendance is _________% he/she can/can\'t take part in sports activity on __/__/____ to__/__/____.'
            para = doc.add_paragraph(content)
            for run in para.runs:
                run.font.size = Inches(14/72)
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            sig_para = doc.add_paragraph('Physical Education Director            Sports Co-ordinator                            Head of the Department')
            for run in sig_para.runs:
                run.font.size = Inches(14/72)
            doc.add_page_break()
    
    elif report_format == 'vtu_bonafide':
        for i in range(5):
            doc.add_paragraph('')
        
        to_para = doc.add_paragraph('To .')
        for run in to_para.runs:
            run.font.size = Inches(14/72)
        
        for i in range(3):
            doc.add_paragraph('')
        
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run('Sub : List of Students participating in ______________(game ) tournament .')
        sub_run.font.size = Inches(14/72)
        sub_run.bold = True
        
        doc.add_paragraph('')
        
        ref_para = doc.add_paragraph('With reference to the above subject , I wish to state that the following Bonafide students of our college will be participating in _____________________________________ tournament .')
        for run in ref_para.runs:
            run.font.size = Inches(14/72)
        
        doc.add_paragraph('')
        
        req_para = doc.add_paragraph('Hence, I request you to kindly permit them and oblige')
        for run in req_para.runs:
            run.font.size = Inches(14/72)
        
        doc.add_paragraph('')
        
        table = doc.add_table(rows=len(students)+1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sl.No.'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'USN'
        hdr_cells[3].text = 'Branch'
        
        for i, student in enumerate(students):
            row_cells = table.rows[i+1].cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = student[1] or ''
            row_cells[2].text = student[6] or ''
            row_cells[3].text = student[5] or ''
        
        doc.add_paragraph('')
        doc.add_paragraph('')
        
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.columns[0].width = Inches(3)
        sig_table.columns[1].width = Inches(3)
        
        left_cell = sig_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.add_run('Physical Education Director\n(for signature)')
        
        right_cell = sig_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_para.add_run('Principal\n(for signature)')
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'complete_{report_format}_report.docx')

if __name__ == '__main__':
    init_db()
    app.run(debug=True)