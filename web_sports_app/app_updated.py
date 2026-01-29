from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, session
import os
import sqlite3
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
import tempfile
import shutil

app = Flask(__name__)
app.secret_key = 'your_secret_key'

UPLOAD_FOLDER = 'static/uploads'
TEMPLATE_FOLDER = 'static/templates'
for folder in [UPLOAD_FOLDER, TEMPLATE_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER

DB_NAME = 'students.db'

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            dob TEXT,
            mother_name TEXT,
            father_name TEXT,
            branch_sem TEXT,
            usn TEXT UNIQUE,
            phone TEXT,
            email TEXT,
            photo_path TEXT,
            sports TEXT,
            blood_group TEXT,
            gender TEXT
        )
    ''')
    try:
        c.execute('ALTER TABLE students ADD COLUMN blood_group TEXT')
        c.execute('ALTER TABLE students ADD COLUMN gender TEXT')
    except:
        pass
    conn.commit()
    conn.close()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/index')
def index():
    return redirect(url_for('home'))

@app.route('/data-entry', methods=['GET', 'POST'])
def data_entry():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()
        blood_group = request.form.get('blood_group', '').strip()
        gender = request.form.get('gender', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))

        # Validation
        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                INSERT INTO students (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports, blood_group, gender)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, blood_group, gender))
            conn.commit()
            conn.close()
            flash('Student saved successfully.', 'success')
            return redirect(url_for('data_entry'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)

    return render_template('data_entry.html')

@app.route('/data-view')
def data_view():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_view.html', students=students)

@app.route('/data-edit')
def data_edit():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_edit.html', students=students)

@app.route('/edit-student/<int:student_id>', methods=['GET', 'POST'])
def edit_student(student_id):
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()
        blood_group = request.form.get('blood_group', '').strip()
        gender = request.form.get('gender', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
        else:
            # Keep existing photo
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('SELECT photo_path FROM students WHERE id = ?', (student_id,))
            result = c.fetchone()
            conn.close()
            if result:
                photo_filename = result[0]

        # Validation
        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                UPDATE students SET name=?, dob=?, mother_name=?, father_name=?, branch_sem=?, usn=?, phone=?, email=?, photo_path=?, sports=?, blood_group=?, gender=?
                WHERE id=?
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, blood_group, gender, student_id))
            conn.commit()
            conn.close()
            flash('Student updated successfully.', 'success')
            return redirect(url_for('data_edit'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)
    
    # GET request - load student data
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students WHERE id = ?', (student_id,))
    student = c.fetchone()
    conn.close()
    
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('data_edit'))
    
    return render_template('edit_student.html', student=student)

@app.route('/delete-student/<int:student_id>', methods=['POST'])
def delete_student(student_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('DELETE FROM students WHERE id = ?', (student_id,))
    conn.commit()
    conn.close()
    flash('Student deleted successfully.', 'success')
    return redirect(url_for('data_edit'))

@app.route('/data-select')
def data_select():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email, sports FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    return render_template('data_select.html', students=students)

@app.route('/template-upload', methods=['GET', 'POST'])
def template_upload():
    if request.method == 'POST':
        template_file = request.files.get('template')
        if template_file and template_file.filename != '':
            template_filename = secure_filename(template_file.filename)
            template_file.save(os.path.join(app.config['TEMPLATE_FOLDER'], template_filename))
            session['template_uploaded'] = True
            flash('Template uploaded successfully.', 'success')
        else:
            flash('Please select a template file.', 'error')
    return render_template('template_upload.html')

@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/edit-report', methods=['POST'])
def edit_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    # Generate default content based on format
    if report_format == 'vtu_eligibility':
        report_title = 'ELIGIBILITYPROFORMA'
        report_content = '''ELIGIBILITYPROFORMA
ELIGIBILITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26
COLLEGE NAME & ADDRESS : DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082
GAME :- ____________
ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________

A | B | C | D | E | F | G
SL NO. | Student Details | Course Details | Academic Details | VTU Previous | Photo | Signature
1 | Name: [NAME]\\nFather: [FATHER]\\nMother: [MOTHER]\\nBranch: [BRANCH]\\nUSN: [USN] | Course: [BRANCH]\\nDuration: 4 Years\\nDOB: [DOB]\\nContact: [PHONE]\\nBlood Group: [BLOOD_GROUP] | PUC Date: ___\\nFirst Admission: ___\\nCurrent Admission: ___ | Game: [SPORTS]\\nYear: ___ | [Photo] | [Signature]'''
    elif report_format == 'hod_bonafide':
        report_title = 'HOD BONAFIDE CERTIFICATE'
        report_content = '''__________________________________________________________________________________
 
:___________________________ ________________________________


This is to certify that Mr/Ms [NAME] is a student of [BRANCH] department studying in _____________ Semester Bearing USN [USN] for academic year 
20__-20__.And his/her present attendance is _________% he/she can/cannot take part in sports activity on __/__/____ to__/__/____.




Physical Education Director            Sports Co-ordinator                            Head of the Department'''
    else:
        report_title = 'VTU BONAFIDE CERTIFICATE'
        report_content = '''Visvesvaraya Technological University
Belagavi - 590018

BONAFIDE CERTIFICATE

Certified that Mr./Ms. [NAME] bearing USN: [USN] is a bonafide student of [College Name] affiliated to Visvesvaraya Technological University.

Course: [BRANCH]
Date of Birth: [DOB]
Father's Name: [FATHER]
Mother's Name: [MOTHER]
Contact: [PHONE]

This certificate is issued for official purposes.

Registrar
Visvesvaraya Technological University'''
    
    return render_template('report_editor.html', 
                         students=students, 
                         selected_students=','.join(selected_ids),
                         report_format=report_format,
                         report_title=report_title,
                         report_content=report_content)

@app.route('/generate-edited-report', methods=['POST'])
def generate_edited_report():
    selected_ids = request.form.get('selected_students').split(',')
    report_title = request.form.get('report_title')
    report_content = request.form.get('report_content')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    for student in students:
        # Replace placeholders
        content = report_content
        content = content.replace('[NAME]', student[1] or '')
        content = content.replace('[USN]', student[6] or '')
        content = content.replace('[DOB]', student[2] or '')
        content = content.replace('[FATHER]', student[4] or '')
        content = content.replace('[MOTHER]', student[3] or '')
        content = content.replace('[BRANCH]', student[5] or '')
        content = content.replace('[PHONE]', student[7] or '')
        content = content.replace('[EMAIL]', student[8] or '')
        content = content.replace('[SPORTS]', student[10] or '')
        content = content.replace('[GENDER]', (student[12] if len(student) > 12 and student[12] else ''))
        content = content.replace('[BLOOD_GROUP]', (student[11] if len(student) > 11 and student[11] else ''))
        
        # Add content to document
        for line in content.split('\\n'):
            doc.add_paragraph(line)
        
        doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name='edited_report.docx')

@app.route('/generate-all-report', methods=['POST'])
def generate_all_report():
    report_format = request.form.get('report_format', 'detailed')

    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students')
    students = c.fetchall()
    conn.close()

    from docx.shared import Inches
    doc = Document()

    if report_format == 'hod_bonafide':
        for student in students:
            # Add logo to right corner
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_run = logo_para.add_run()

            logo_path = r'C:\Users\user\Desktop\dsatm project\web_sports_app\static\image1.png'
            if os.path.exists(logo_path):
                logo_run.add_picture(logo_path, width=Inches(1.5))

            doc.add_paragraph('')
            doc.add_paragraph('')

            content = f'This is to certify that Mr/Ms {student[1] or "____________________________"} is a student of {student[5] or "___________________________"} department studying in _____________ Semester Bearing USN {student[6] or "_____________________________"} for academic year \n20__-20__.And his/her present attendance is _________% he/she can/cannot take part in sports activity on __/__/____ to__/__/____.'
            doc.add_paragraph(content)

            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')

            doc.add_paragraph('Physical Education Director            Sports Co-ordinator                            Head of the Department')
            doc.add_page_break()
    elif report_format == 'vtu_eligibility':
        for student in students:
            doc.add_paragraph('ELIGIBILITYPROFORMA')
            doc.add_paragraph('ELIGIBILITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26')
            doc.add_paragraph('COLLEGE NAME & ADDRESS : DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082')
            doc.add_paragraph('GAME :- ____________')
            doc.add_paragraph('ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________')
            doc.add_paragraph('')

            table = doc.add_table(rows=2, cols=7)
            table.style = 'Table Grid'

            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'SL NO.'
            hdr_cells[1].text = 'Student Details'
            hdr_cells[2].text = 'Course Details'
            hdr_cells[3].text = 'Academic Details'
            hdr_cells[4].text = 'VTU Previous'
            hdr_cells[5].text = 'Photo'
            hdr_cells[6].text = 'Signature'

            # Data
            row_cells = table.rows[1].cells
            row_cells[0].text = '1'
            row_cells[1].text = f'Name: {student[1] or ""}\nFather: {student[4] or ""}\nMother: {student[3] or ""}\nBranch: {student[5] or ""}\nUSN: {student[6] or ""}'
            row_cells[2].text = f'Course: {student[5] or ""}\nDuration: 4 Years\nDOB: {student[2] or ""}\nContact: {student[7] or ""}\nBlood Group: {student[11] or ""}'
            row_cells[3].text = 'PUC Date: ___\nFirst Admission: ___\nCurrent Admission: ___'
            row_cells[4].text = f'Game: {student[10] or ""}\nYear: ___'
            if student[9]:
                photo_path = os.path.join(app.config['UPLOAD_FOLDER'], student[9])
                if os.path.exists(photo_path):
                    para = row_cells[5].add_paragraph()
                    run = para.add_run()
                    run.add_picture(photo_path, width=Inches(1))
                else:
                    row_cells[5].text = '[Photo]'
            else:
                row_cells[5].text = '[Photo]'
            row_cells[6].text = '[Signature]'

            doc.add_page_break()
    elif report_format == 'vtu_bonafide':
        for student in students:
            content = f'''Visvesvaraya Technological University
Belagavi - 590018

BONAFIDE CERTIFICATE

Certified that Mr./Ms. {student[1] or ''} bearing USN: {student[6] or ''} is a bonafide student of [College Name] affiliated to Visvesvaraya Technological University.

Course: {student[5] or ''}
Date of Birth: {student[2] or ''}
Father's Name: {student[4] or ''}
Mother's Name: {student[3] or ''}
Contact: {student[7] or ''}

This certificate is issued for official purposes.

Registrar
Visvesvaraya Technological University'''
            for line in content.split('\n'):
                doc.add_paragraph(line)
            doc.add_page_break()

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    return send_file(temp_file.name, as_attachment=True, download_name=f'{report_format}_all_report.docx')

@app.route('/generate-report', methods=['POST'])
def generate_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')

    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))

    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()

    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()

    from docx.shared import Inches
    doc = Document()

    if report_format == 'hod_bonafide':
        for student in students:
            # Add logo to right corner
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_run = logo_para.add_run()

            logo_path = r'C:\Users\user\Desktop\dsatm project\web_sports_app\static\image1.png'
            if os.path.exists(logo_path):
                logo_run.add_picture(logo_path, width=Inches(1.5))

            doc.add_paragraph('')
            doc.add_paragraph('')

            content = f'This is to certify that Mr/Ms {student[1] or "____________________________"} is a student of {student[5] or "___________________________"} department studying in _____________ Semester Bearing USN {student[6] or "_____________________________"} for academic year \n20__-20__.And his/her present attendance is _________% he/she can/cannot take part in sports activity on __/__/____ to__/__/____.'
            doc.add_paragraph(content)

            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')

            doc.add_paragraph('Physical Education Director            Sports Co-ordinator                            Head of the Department')
            doc.add_page_break()
    elif report_format == 'vtu_eligibility':
        for student in students:
            doc.add_paragraph('ELIGIBILITYPROFORMA')
            doc.add_paragraph('ELIGIBILITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26')
            doc.add_paragraph('COLLEGE NAME & ADDRESS : DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082')
            doc.add_paragraph('GAME :- ____________')
            doc.add_paragraph('ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________')
            doc.add_paragraph('')

            table = doc.add_table(rows=2, cols=7)
            table.style = 'Table Grid'

            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'SL NO.'
            hdr_cells[1].text = 'Student Details'
            hdr_cells[2].text = 'Course Details'
            hdr_cells[3].text = 'Academic Details'
            hdr_cells[4].text = 'VTU Previous'
            hdr_cells[5].text = 'Photo'
            hdr_cells[6].text = 'Signature'

            # Data
            row_cells = table.rows[1].cells
            row_cells[0].text = '1'
            row_cells[1].text = f'Name: {student[1] or ""}\nFather: {student[4] or ""}\nMother: {student[3] or ""}\nBranch: {student[5] or ""}\nUSN: {student[6] or ""}'
            row_cells[2].text = f'Course: {student[5] or ""}\nDuration: 4 Years\nDOB: {student[2] or ""}\nContact: {student[7] or ""}\nBlood Group: {student[11] or ""}'
            row_cells[3].text = 'PUC Date: ___\nFirst Admission: ___\nCurrent Admission: ___'
            row_cells[4].text = f'Game: {student[10] or ""}\nYear: ___'
            if student[9]:
                photo_path = os.path.join(app.config['UPLOAD_FOLDER'], student[9])
                if os.path.exists(photo_path):
                    para = row_cells[5].add_paragraph()
                    run = para.add_run()
                    run.add_picture(photo_path, width=Inches(1))
                else:
                    row_cells[5].text = '[Photo]'
            else:
                row_cells[5].text = '[Photo]'
            row_cells[6].text = '[Signature]'

            doc.add_page_break()
    elif report_format == 'vtu_bonafide':
        for student in students:
            content = f'''Visvesvaraya Technological University
Belagavi - 590018

BONAFIDE CERTIFICATE

Certified that Mr./Ms. {student[1] or ''} bearing USN: {student[6] or ''} is a bonafide student of [College Name] affiliated to Visvesvaraya Technological University.

Course: {student[5] or ''}
Date of Birth: {student[2] or ''}
Father's Name: {student[4] or ''}
Mother's Name: {student[3] or ''}
Contact: {student[7] or ''}

This certificate is issued for official purposes.

Registrar
Visvesvaraya Technological University'''
            for line in content.split('\n'):
                doc.add_paragraph(line)
            doc.add_page_break()

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    return send_file(temp_file.name, as_attachment=True, download_name=f'{report_format}_report.docx')

if __name__ == '__main__':
    init_db()
    app.run(debug=True)