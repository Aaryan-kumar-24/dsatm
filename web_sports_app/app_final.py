from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, session
import os
import sqlite3
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import tempfile
import shutil

app = Flask(__name__)
app.secret_key = 'your_secret_key'

UPLOAD_FOLDER = 'static/uploads'
TEMPLATE_FOLDER = 'static/templates'
for folder in [UPLOAD_FOLDER, TEMPLATE_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER

DB_NAME = 'students.db'

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            dob TEXT,
            mother_name TEXT,
            father_name TEXT,
            branch_sem TEXT,
            usn TEXT UNIQUE,
            phone TEXT,
            email TEXT,
            photo_path TEXT,
            sports TEXT,
            blood_group TEXT,
            gender TEXT
        )
    ''')
    try:
        c.execute('ALTER TABLE students ADD COLUMN blood_group TEXT')
        c.execute('ALTER TABLE students ADD COLUMN gender TEXT')
    except:
        pass
    conn.commit()
    conn.close()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/data-entry', methods=['GET', 'POST'])
def data_entry():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()
        blood_group = request.form.get('blood_group', '').strip()
        gender = request.form.get('gender', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))

        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                INSERT INTO students (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports, blood_group, gender)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, blood_group, gender))
            conn.commit()
            conn.close()
            flash('Student saved successfully.', 'success')
            return redirect(url_for('data_entry'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)

    return render_template('data_entry.html')

@app.route('/data-view')
def data_view():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_view.html', students=students)

@app.route('/data-select')
def data_select():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email, sports FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    return render_template('data_select.html', students=students)

@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/generate-report', methods=['POST'])
def generate_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    if report_format == 'hod_bonafide':
        for student in students:
            content = f'This is to certify that Mr/Ms {student[1] or "____________________________"} is a student of {student[5] or "___________________________"} department studying in _____________ Semester Bearing USN {student[6] or "_____________________________"} for academic year \n20__-20__.And his/her present attendance is _________% he/she can/can\'t take part in sports activity on __/__/____ to__/__/____.'
            doc.add_paragraph(content)
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            doc.add_paragraph('Physical Education Director            Sports Co-ordinator                            Head of the Department')
            doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'{report_format}_report.docx')

if __name__ == '__main__':
    init_db()
    app.run(debug=True)