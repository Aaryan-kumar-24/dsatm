from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify
import os
import sqlite3
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
import tempfile

app = Flask(__name__)
app.secret_key = 'your_secret_key'

UPLOAD_FOLDER = 'static/uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

DB_NAME = 'students.db'

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            dob TEXT,
            mother_name TEXT,
            father_name TEXT,
            branch_sem TEXT,
            usn TEXT UNIQUE,
            phone TEXT,
            email TEXT,
            photo_path TEXT,
            sports TEXT
        )
    ''')
    conn.commit()
    conn.close()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/index')
def index():
    return redirect(url_for('home'))

@app.route('/data-entry', methods=['GET', 'POST'])
def data_entry():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))

        # Validation
        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                INSERT INTO students (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports))
            conn.commit()
            conn.close()
            flash('Student saved successfully.', 'success')
            return redirect(url_for('data_entry'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)

    return render_template('data_entry.html')

@app.route('/data-view')
def data_view():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_view.html', students=students)

@app.route('/data-edit')
def data_edit():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_edit.html', students=students)

@app.route('/edit-student/<int:student_id>', methods=['GET', 'POST'])
def edit_student(student_id):
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
        else:
            # Keep existing photo
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('SELECT photo_path FROM students WHERE id = ?', (student_id,))
            result = c.fetchone()
            conn.close()
            if result:
                photo_filename = result[0]

        # Validation
        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                UPDATE students SET name=?, dob=?, mother_name=?, father_name=?, branch_sem=?, usn=?, phone=?, email=?, photo_path=?, sports=?
                WHERE id=?
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, student_id))
            conn.commit()
            conn.close()
            flash('Student updated successfully.', 'success')
            return redirect(url_for('data_edit'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)
    
    # GET request - load student data
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students WHERE id = ?', (student_id,))
    student = c.fetchone()
    conn.close()
    
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('data_edit'))
    
    return render_template('edit_student.html', student=student)

@app.route('/delete-student/<int:student_id>', methods=['POST'])
def delete_student(student_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('DELETE FROM students WHERE id = ?', (student_id,))
    conn.commit()
    conn.close()
    flash('Student deleted successfully.', 'success')
    return redirect(url_for('data_edit'))

@app.route('/data-select')
def data_select():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email, sports FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    return render_template('data_select.html', students=students)

@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/generate-report', methods=['POST'])
def generate_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    if report_format == 'vtu_eligibility':
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        title = doc.add_paragraph()
        run = title.add_run('ELIGIBILITYPROFORMA')
        run.bold = True
        run.underline = True
        
        subtitle = doc.add_paragraph()
        subtitle.add_run('ELIGIBITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26')
        
        college_info = doc.add_paragraph()
        college_info.add_run('COLLEGE NAME & ADDRESS : ')
        run = college_info.add_run('DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082')
        run.underline = True
        
        game = doc.add_paragraph()
        game.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        game.add_run('GAME :- ____________')
        
        org_college = doc.add_paragraph()
        org_college.add_run('ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________')
        
        doc.add_paragraph('')
        
        table = doc.add_table(rows=len(students)+2, cols=7)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'A'
        hdr_cells[1].text = '        B'
        hdr_cells[2].text = '           C'
        hdr_cells[3].text = '        D'
        hdr_cells[4].text = '          E'
        hdr_cells[5].text = '          F'
        hdr_cells[6].text = '         G'
        
        hdr_cells2 = table.rows[1].cells
        hdr_cells2[0].text = 'SL\nNO.'
        hdr_cells2[1].text = 'A)Name of the student \nB)Name of father \nC)Name of mother \nD)Semester/Branch\nE)University seat No.(USN)'
        hdr_cells2[2].text = 'A) Name of the course \nB) Duration of the course\nC)Date of birth \nD)Contact No.\nE) Blood Group'
        hdr_cells2[3].text = 'A)Passing Date of PUC/DIPLOMA\nB) Date of first admission to present course\nC)Date of admission to present Class / Semester \n'
        hdr_cells2[4].text = 'VTU Previous Preresentation \na)Game and Year'
        hdr_cells2[5].text = 'Photo with principle attestation'
        hdr_cells2[6].text = 'Signature of the student '
        
        for i, student in enumerate(students):
            row_cells = table.rows[i+2].cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = f'a) {student[1] or "[Name]"}\nb) {student[4] or "[Father Name]"}\nc) {student[3] or "[Mother Name]"}\nd) {student[5] or "[Branch/Semester]"}\ne) {student[6] or "[USN]"}'
            row_cells[2].text = f'A) {student[5] or "[Course]"}\nB) 4 Years\nC) {student[2] or "[DOB]"}\nD) {student[7] or "[Contact]"}\nE) [Blood Group]'
            row_cells[3].text = 'A) [PUC Date]\nB) [First Admission]\nC) [Current Admission]\n\n'
            row_cells[4].text = f'A) {student[10] or "[Sports/Year]"}\nB)\nC)\nD)\nE)'
            row_cells[5].text = ''
            row_cells[6].text = ''
    
    elif report_format == 'hod_bonafide':
        doc.add_heading('HOD BONAFIDE CERTIFICATE', 0)
        doc.add_paragraph('[College Name]')
        doc.add_paragraph('[College Address]')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph(f'This is to certify that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student of our department studying in {student[5] or "[Branch & Semester]"}.')
            doc.add_paragraph('')
            doc.add_paragraph(f'Personal Details:')
            doc.add_paragraph(f'Date of Birth: {student[2] or "[DOB]"}')
            doc.add_paragraph(f'Father\'s Name: {student[4] or "[Father Name]"}')
            doc.add_paragraph(f'Contact: {student[7] or "[Phone]"}')
            doc.add_paragraph(f'Sports: {student[10] or "[Sports Activities]"}')
            doc.add_paragraph('')
            doc.add_paragraph('This certificate is issued for sports participation purposes.')
            doc.add_paragraph('')
            doc.add_paragraph('Head of Department')
            doc.add_paragraph('[Department Name]')
            doc.add_page_break()
    
    elif report_format == 'vtu_bonafide':
        doc.add_heading('VTU BONAFIDE CERTIFICATE', 0)
        doc.add_paragraph('Visvesvaraya Technological University')
        doc.add_paragraph('Belagavi - 590018')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph(f'Certified that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student of [College Name] affiliated to Visvesvaraya Technological University.')
            doc.add_paragraph('')
            doc.add_paragraph(f'Course: {student[5] or "[Branch & Semester]"}')
            doc.add_paragraph(f'Date of Birth: {student[2] or "[DOB]"}')
            doc.add_paragraph(f'Father\'s Name: {student[4] or "[Father Name]"}')
            doc.add_paragraph(f'Mother\'s Name: {student[3] or "[Mother Name]"}')
            doc.add_paragraph(f'Contact: {student[7] or "[Phone]"}')
            doc.add_paragraph('')
            doc.add_paragraph('This certificate is issued for official purposes.')
            doc.add_paragraph('')
            doc.add_paragraph('Registrar')
            doc.add_paragraph('Visvesvaraya Technological University')
            doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'{report_format}_report.docx')

@app.route('/generate-all-report', methods=['POST'])
def generate_all_report():
    report_format = request.form.get('report_format', 'detailed')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    
    if not students:
        flash('No students found to generate report.', 'error')
        return redirect(url_for('report'))
    
    doc = Document()
    
    if report_format == 'vtu_eligibility':
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        title = doc.add_paragraph()
        run = title.add_run('ELIGIBILITYPROFORMA')
        run.bold = True
        run.underline = True
        
        subtitle = doc.add_paragraph()
        subtitle.add_run('ELIGIBITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26')
        
        college_info = doc.add_paragraph()
        college_info.add_run('COLLEGE NAME & ADDRESS : ')
        run = college_info.add_run('DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082')
        run.underline = True
        
        game = doc.add_paragraph()
        game.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        game.add_run('GAME :- ____________')
        
        org_college = doc.add_paragraph()
        org_college.add_run('ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________')
        
        doc.add_paragraph('')
        
        table = doc.add_table(rows=len(students)+2, cols=7)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'A'
        hdr_cells[1].text = '        B'
        hdr_cells[2].text = '           C'
        hdr_cells[3].text = '        D'
        hdr_cells[4].text = '          E'
        hdr_cells[5].text = '          F'
        hdr_cells[6].text = '         G'
        
        hdr_cells2 = table.rows[1].cells
        hdr_cells2[0].text = 'SL\nNO.'
        hdr_cells2[1].text = 'A)Name of the student \nB)Name of father \nC)Name of mother \nD)Semester/Branch\nE)University seat No.(USN)'
        hdr_cells2[2].text = 'A) Name of the course \nB) Duration of the course\nC)Date of birth \nD)Contact No.\nE) Blood Group'
        hdr_cells2[3].text = 'A)Passing Date of PUC/DIPLOMA\nB) Date of first admission to present course\nC)Date of admission to present Class / Semester \n'
        hdr_cells2[4].text = 'VTU Previous Preresentation \na)Game and Year'
        hdr_cells2[5].text = 'Photo with principle attestation'
        hdr_cells2[6].text = 'Signature of the student '
        
        for i, student in enumerate(students):
            row_cells = table.rows[i+2].cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = f'a) {student[1] or "[Name]"}\nb) {student[4] or "[Father Name]"}\nc) {student[3] or "[Mother Name]"}\nd) {student[5] or "[Branch/Semester]"}\ne) {student[6] or "[USN]"}'
            row_cells[2].text = f'A) {student[5] or "[Course]"}\nB) 4 Years\nC) {student[2] or "[DOB]"}\nD) {student[7] or "[Contact]"}\nE) [Blood Group]'
            row_cells[3].text = 'A) [PUC Date]\nB) [First Admission]\nC) [Current Admission]\n\n'
            row_cells[4].text = f'A) {student[10] or "[Sports/Year]"}\nB)\nC)\nD)\nE)'
            row_cells[5].text = ''
            row_cells[6].text = ''
    
    elif report_format == 'hod_bonafide':
        doc.add_heading('HOD BONAFIDE CERTIFICATES', 0)
        doc.add_paragraph(f'Total Students: {len(students)}')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('HOD BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph('[College Name]')
            doc.add_paragraph('')
            doc.add_paragraph(f'This is to certify that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student of our department.')
            doc.add_paragraph(f'Course: {student[5] or "[Branch & Semester]"}')
            doc.add_paragraph(f'Contact: {student[7] or "[Phone]"}')
            doc.add_paragraph('')
            doc.add_paragraph('Head of Department')
            doc.add_paragraph('[Department Name]')
            doc.add_page_break()
    
    elif report_format == 'vtu_bonafide':
        doc.add_heading('VTU BONAFIDE CERTIFICATES', 0)
        doc.add_paragraph(f'Total Students: {len(students)}')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('VTU BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph('Visvesvaraya Technological University')
            doc.add_paragraph('Belagavi - 590018')
            doc.add_paragraph('')
            doc.add_paragraph(f'Certified that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student affiliated to VTU.')
            doc.add_paragraph(f'Course: {student[5] or "[Branch & Semester]"}')
            doc.add_paragraph('')
            doc.add_paragraph('Registrar')
            doc.add_paragraph('Visvesvaraya Technological University')
            doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'complete_{report_format}_report.docx')

if __name__ == '__main__':
    init_db()
    app.run(debug=True)