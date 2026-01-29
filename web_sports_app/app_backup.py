from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, session
import os
import sqlite3
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
import tempfile
import shutil

app = Flask(__name__)
app.secret_key = 'your_secret_key'

UPLOAD_FOLDER = 'static/uploads'
TEMPLATE_FOLDER = 'static/templates'
for folder in [UPLOAD_FOLDER, TEMPLATE_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER

DB_NAME = 'students.db'

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            dob TEXT,
            mother_name TEXT,
            father_name TEXT,
            branch_sem TEXT,
            usn TEXT UNIQUE,
            phone TEXT,
            email TEXT,
            photo_path TEXT,
            sports TEXT,
            blood_group TEXT,
            gender TEXT
        )
    ''')
    try:
        c.execute('ALTER TABLE students ADD COLUMN blood_group TEXT')
        c.execute('ALTER TABLE students ADD COLUMN gender TEXT')
    except:
        pass
    conn.commit()
    conn.close()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/index')
def index():
    return redirect(url_for('home'))

@app.route('/data-entry', methods=['GET', 'POST'])
def data_entry():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))

        # Validation
        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                INSERT INTO students (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports))
            conn.commit()
            conn.close()
            flash('Student saved successfully.', 'success')
            return redirect(url_for('data_entry'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)

    return render_template('data_entry.html')

@app.route('/data-view')
def data_view():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_view.html', students=students)

@app.route('/data-edit')
def data_edit():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email FROM students ORDER BY id DESC')
    students = c.fetchall()
    conn.close()
    return render_template('data_edit.html', students=students)

@app.route('/edit-student/<int:student_id>', methods=['GET', 'POST'])
def edit_student(student_id):
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        dob = request.form.get('dob', '').strip()
        mother_name = request.form.get('mother_name', '').strip()
        father_name = request.form.get('father_name', '').strip()
        branch_sem = request.form.get('branch_sem', '').strip()
        usn = request.form.get('usn', '').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form.get('email', '').strip()
        sports = request.form.get('sports', '').strip()

        photo = request.files.get('photo')
        photo_filename = None
        if photo and photo.filename != '':
            photo_filename = secure_filename(photo.filename)
            photo.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
        else:
            # Keep existing photo
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('SELECT photo_path FROM students WHERE id = ?', (student_id,))
            result = c.fetchone()
            conn.close()
            if result:
                photo_filename = result[0]

        # Validation
        if not name:
            flash('Name is required.', 'error')
            return redirect(request.url)
        if len(usn) != 10:
            flash('USN must be exactly 10 characters.', 'error')
            return redirect(request.url)
        if not (phone.isdigit() and len(phone) == 10):
            flash('Phone number must be numeric and exactly 10 digits.', 'error')
            return redirect(request.url)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                UPDATE students SET name=?, dob=?, mother_name=?, father_name=?, branch_sem=?, usn=?, phone=?, email=?, photo_path=?, sports=?
                WHERE id=?
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_filename, sports, student_id))
            conn.commit()
            conn.close()
            flash('Student updated successfully.', 'success')
            return redirect(url_for('data_edit'))
        except sqlite3.IntegrityError:
            flash('USN must be unique. A student with this USN already exists.', 'error')
            return redirect(request.url)
    
    # GET request - load student data
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students WHERE id = ?', (student_id,))
    student = c.fetchone()
    conn.close()
    
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('data_edit'))
    
    return render_template('edit_student.html', student=student)

@app.route('/delete-student/<int:student_id>', methods=['POST'])
def delete_student(student_id):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('DELETE FROM students WHERE id = ?', (student_id,))
    conn.commit()
    conn.close()
    flash('Student deleted successfully.', 'success')
    return redirect(url_for('data_edit'))

@app.route('/data-select')
def data_select():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn, phone, email, sports FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    return render_template('data_select.html', students=students)

@app.route('/report')
def report():
    return render_template('report.html')

@app.route('/edit-report', methods=['POST'])
def edit_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    # Generate default content based on format
    if report_format == 'vtu_eligibility':
        report_title = 'ELIGIBILITYPROFORMA'
        report_content = '''ELIGIBILITYPROFORMA
ELIGIBITY PROFORMA OF PLAYERS REPRESENTING COLLEGE IN VTU INTER-COLLEGIATE SPORTS/TOURNAMENT 2025-26
COLLEGE NAME & ADDRESS : DAYANANDA SAGAR ACADEMY OF TECHNOLOGY AND MANAGEMENT , BANGALURU 560082
GAME :- ____________
ORGANISING COLLEGE:- __________________________________DIVISION : Bangalore division _________________

A | B | C | D | E | F | G
SL NO. | Student Details | Course Details | Academic Details | VTU Previous | Photo | Signature
1 | Name: [NAME]\nFather: [FATHER]\nMother: [MOTHER]\nBranch: [BRANCH]\nUSN: [USN] | Course: [BRANCH]\nDuration: 4 Years\nDOB: [DOB]\nContact: [PHONE]\nBlood Group: [BLOOD_GROUP] | PUC Date: ___\nFirst Admission: ___\nCurrent Admission: ___ | Game: [SPORTS]\nYear: ___ | [Photo] | [Signature]'''
    elif report_format == 'hod_bonafide':
        report_title = 'HOD BONAFIDE CERTIFICATE'
        report_content = '''[College Name]
[College Address]

BONAFIDE CERTIFICATE

This is to certify that Mr./Ms. [NAME] bearing USN: [USN] is a bonafide student of our department studying in [BRANCH].

Personal Details:
Date of Birth: [DOB]
Father's Name: [FATHER]
Gender: [GENDER]
Contact: [PHONE]
Sports: [SPORTS]

This certificate is issued for sports participation purposes.

Head of Department
[Department Name]'''
    else:
        report_title = 'VTU BONAFIDE CERTIFICATE'
        report_content = '''Visvesvaraya Technological University
Belagavi - 590018

BONAFIDE CERTIFICATE

Certified that Mr./Ms. [NAME] bearing USN: [USN] is a bonafide student of [College Name] affiliated to Visvesvaraya Technological University.

Course: [BRANCH]
Date of Birth: [DOB]
Father's Name: [FATHER]
Mother's Name: [MOTHER]
Contact: [PHONE]

This certificate is issued for official purposes.

Registrar
Visvesvaraya Technological University'''
    
    return render_template('report_editor.html', 
                         students=students, 
                         selected_students=','.join(selected_ids),
                         report_format=report_format,
                         report_title=report_title,
                         report_content=report_content)

@app.route('/generate-edited-report', methods=['POST'])
def generate_edited_report():
    selected_ids = request.form.get('selected_students').split(',')
    report_title = request.form.get('report_title')
    report_content = request.form.get('report_content')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    for student in students:
        # Replace placeholders
        content = report_content
        content = content.replace('[NAME]', student[1] or '')
        content = content.replace('[USN]', student[6] or '')
        content = content.replace('[DOB]', student[2] or '')
        content = content.replace('[FATHER]', student[4] or '')
        content = content.replace('[MOTHER]', student[3] or '')
        content = content.replace('[BRANCH]', student[5] or '')
        content = content.replace('[PHONE]', student[7] or '')
        content = content.replace('[EMAIL]', student[8] or '')
        content = content.replace('[SPORTS]', student[10] or '')
        content = content.replace('[GENDER]', (student[12] if len(student) > 12 and student[12] else ''))
        content = content.replace('[BLOOD_GROUP]', (student[11] if len(student) > 11 and student[11] else ''))
        
        # Add content to document
        for line in content.split('\n'):
            doc.add_paragraph(line)
        
        doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name='edited_report.docx')

@app.route('/generate-report', methods=['POST'])
def generate_report():
    selected_ids = request.form.getlist('selected_students')
    report_format = request.form.get('report_format', 'detailed')
    
    if not selected_ids:
        flash('Please select at least one student.', 'error')
        return redirect(url_for('data_select'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document()
    
    if report_format == 'vtu_eligibility':
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for student in students:
            # Header
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run('VISVESVARAYA TECHNOLOGICAL UNIVERSITY')
            run.bold = True
            run.font.size = 140000  # 14pt
            
            subheader = doc.add_paragraph()
            subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subheader.add_run('BELAGAVI - 590018')
            run.font.size = 120000  # 12pt
            
            doc.add_paragraph('')
            
            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run('ELIGIBILITY CERTIFICATE')
            run.bold = True
            run.underline = True
            run.font.size = 140000
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            # Certificate content
            content = doc.add_paragraph()
            content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content.add_run(f'\t\tThis is to certify that Mr./Ms. ').font.size = 120000
            name_run = content.add_run(f'{student[1]}')
            name_run.bold = True
            name_run.font.size = 120000
            content.add_run(f' bearing USN: ').font.size = 120000
            usn_run = content.add_run(f'{student[6]}')
            usn_run.bold = True
            usn_run.font.size = 120000
            content.add_run(f' is a bonafide student of our institution studying in ').font.size = 120000
            branch_run = content.add_run(f'{student[5] or "[Branch & Semester]"}')
            branch_run.bold = True
            branch_run.font.size = 120000
            content.add_run('.').font.size = 120000
            
            doc.add_paragraph('')
            
            # Personal details
            details = doc.add_paragraph()
            details.alignment = WD_ALIGN_PARAGRAPH.LEFT
            details.add_run(f'Date of Birth\t\t: {student[2] or "[DOB]"}').font.size = 120000
            
            father = doc.add_paragraph()
            father.alignment = WD_ALIGN_PARAGRAPH.LEFT
            father.add_run(f'Father\'s Name\t\t: {student[4] or "[Father Name]"}').font.size = 120000
            
            mother = doc.add_paragraph()
            mother.alignment = WD_ALIGN_PARAGRAPH.LEFT
            mother.add_run(f'Mother\'s Name\t\t: {student[3] or "[Mother Name]"}').font.size = 120000
            
            sports = doc.add_paragraph()
            sports.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sports.add_run(f'Sports Activity\t\t: {student[10] or "[Sports Activities]"}').font.size = 120000
            
            doc.add_paragraph('')
            
            # Eligibility statement
            eligibility = doc.add_paragraph()
            eligibility.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            eligibility.add_run('\t\tHe/She is eligible to participate in Inter-Collegiate/University level sports competitions as per VTU guidelines and regulations.').font.size = 120000
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            # Signature section
            signature_table = doc.add_table(rows=1, cols=2)
            signature_table.autofit = False
            signature_table.columns[0].width = Inches(3)
            signature_table.columns[1].width = Inches(3)
            
            left_cell = signature_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left_para.add_run('Date: _______________').font.size = 120000
            
            right_cell = signature_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para.add_run('Principal\n[College Name]\n[College Address]').font.size = 120000
            
            doc.add_page_break()
    
    elif report_format == 'hod_bonafide':
        doc.add_heading('HOD BONAFIDE CERTIFICATE', 0)
        doc.add_paragraph('[College Name]')
        doc.add_paragraph('[College Address]')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph(f'This is to certify that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student of our department studying in {student[5] or "[Branch & Semester]"}.')
            doc.add_paragraph('')
            doc.add_paragraph(f'Personal Details:')
            doc.add_paragraph(f'Date of Birth: {student[2] or "[DOB]"}')
            doc.add_paragraph(f'Father\'s Name: {student[4] or "[Father Name]"}')
            doc.add_paragraph(f'Contact: {student[7] or "[Phone]"}')
            doc.add_paragraph(f'Sports: {student[10] or "[Sports Activities]"}')
            doc.add_paragraph('')
            doc.add_paragraph('This certificate is issued for sports participation purposes.')
            doc.add_paragraph('')
            doc.add_paragraph('Head of Department')
            doc.add_paragraph('[Department Name]')
            doc.add_page_break()
    
    elif report_format == 'vtu_bonafide':
        doc.add_heading('VTU BONAFIDE CERTIFICATE', 0)
        doc.add_paragraph('Visvesvaraya Technological University')
        doc.add_paragraph('Belagavi - 590018')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph(f'Certified that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student of [College Name] affiliated to Visvesvaraya Technological University.')
            doc.add_paragraph('')
            doc.add_paragraph(f'Course: {student[5] or "[Branch & Semester]"}')
            doc.add_paragraph(f'Date of Birth: {student[2] or "[DOB]"}')
            doc.add_paragraph(f'Father\'s Name: {student[4] or "[Father Name]"}')
            doc.add_paragraph(f'Mother\'s Name: {student[3] or "[Mother Name]"}')
            doc.add_paragraph(f'Contact: {student[7] or "[Phone]"}')
            doc.add_paragraph('')
            doc.add_paragraph('This certificate is issued for official purposes.')
            doc.add_paragraph('')
            doc.add_paragraph('Registrar')
            doc.add_paragraph('Visvesvaraya Technological University')
            doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'{report_format}_report.docx')

@app.route('/generate-all-report', methods=['POST'])
def generate_all_report():
    report_format = request.form.get('report_format', 'detailed')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT * FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    
    if not students:
        flash('No students found to generate report.', 'error')
        return redirect(url_for('report'))
    
    doc = Document()
    
    if report_format == 'vtu_eligibility':
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for student in students:
            # Header
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run('VISVESVARAYA TECHNOLOGICAL UNIVERSITY')
            run.bold = True
            run.font.size = 140000  # 14pt
            
            subheader = doc.add_paragraph()
            subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subheader.add_run('BELAGAVI - 590018')
            run.font.size = 120000  # 12pt
            
            doc.add_paragraph('')
            
            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run('ELIGIBILITY CERTIFICATE')
            run.bold = True
            run.underline = True
            run.font.size = 140000
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            # Certificate content
            content = doc.add_paragraph()
            content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content.add_run(f'\t\tThis is to certify that Mr./Ms. ').font.size = 120000
            name_run = content.add_run(f'{student[1]}')
            name_run.bold = True
            name_run.font.size = 120000
            content.add_run(f' bearing USN: ').font.size = 120000
            usn_run = content.add_run(f'{student[6]}')
            usn_run.bold = True
            usn_run.font.size = 120000
            content.add_run(f' is a bonafide student of our institution studying in ').font.size = 120000
            branch_run = content.add_run(f'{student[5] or "[Branch & Semester]"}')
            branch_run.bold = True
            branch_run.font.size = 120000
            content.add_run('.').font.size = 120000
            
            doc.add_paragraph('')
            
            # Personal details
            details = doc.add_paragraph()
            details.alignment = WD_ALIGN_PARAGRAPH.LEFT
            details.add_run(f'Date of Birth\t\t: {student[2] or "[DOB]"}').font.size = 120000
            
            father = doc.add_paragraph()
            father.alignment = WD_ALIGN_PARAGRAPH.LEFT
            father.add_run(f'Father\'s Name\t\t: {student[4] or "[Father Name]"}').font.size = 120000
            
            mother = doc.add_paragraph()
            mother.alignment = WD_ALIGN_PARAGRAPH.LEFT
            mother.add_run(f'Mother\'s Name\t\t: {student[3] or "[Mother Name]"}').font.size = 120000
            
            sports = doc.add_paragraph()
            sports.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sports.add_run(f'Sports Activity\t\t: {student[10] or "[Sports Activities]"}').font.size = 120000
            
            doc.add_paragraph('')
            
            # Eligibility statement
            eligibility = doc.add_paragraph()
            eligibility.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            eligibility.add_run('\t\tHe/She is eligible to participate in Inter-Collegiate/University level sports competitions as per VTU guidelines and regulations.').font.size = 120000
            
            doc.add_paragraph('')
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            # Signature section
            signature_table = doc.add_table(rows=1, cols=2)
            signature_table.autofit = False
            signature_table.columns[0].width = Inches(3)
            signature_table.columns[1].width = Inches(3)
            
            left_cell = signature_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left_para.add_run('Date: _______________').font.size = 120000
            
            right_cell = signature_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para.add_run('Principal\n[College Name]\n[College Address]').font.size = 120000
            
            doc.add_page_break()
    
    elif report_format == 'hod_bonafide':
        doc.add_heading('HOD BONAFIDE CERTIFICATES', 0)
        doc.add_paragraph(f'Total Students: {len(students)}')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('HOD BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph('[College Name]')
            doc.add_paragraph('')
            doc.add_paragraph(f'This is to certify that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student of our department.')
            doc.add_paragraph(f'Course: {student[5] or "[Branch & Semester]"}')
            doc.add_paragraph(f'Contact: {student[7] or "[Phone]"}')
            doc.add_paragraph('')
            doc.add_paragraph('Head of Department')
            doc.add_paragraph('[Department Name]')
            doc.add_page_break()
    
    elif report_format == 'vtu_bonafide':
        doc.add_heading('VTU BONAFIDE CERTIFICATES', 0)
        doc.add_paragraph(f'Total Students: {len(students)}')
        doc.add_paragraph('')
        
        for student in students:
            doc.add_heading('VTU BONAFIDE CERTIFICATE', level=1)
            doc.add_paragraph('Visvesvaraya Technological University')
            doc.add_paragraph('Belagavi - 590018')
            doc.add_paragraph('')
            doc.add_paragraph(f'Certified that Mr./Ms. {student[1]} bearing USN: {student[6]} is a bonafide student affiliated to VTU.')
            doc.add_paragraph(f'Course: {student[5] or "[Branch & Semester]"}')
            doc.add_paragraph('')
            doc.add_paragraph('Registrar')
            doc.add_paragraph('Visvesvaraya Technological University')
            doc.add_page_break()
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(temp_file.name, as_attachment=True, download_name=f'complete_{report_format}_report.docx')

@app.route('/template-upload', methods=['GET', 'POST'])
def template_upload():
    if request.method == 'POST':
        template = request.files.get('template')
        if template and template.filename.endswith('.docx'):
            filename = secure_filename(template.filename)
            template_path = os.path.join(app.config['TEMPLATE_FOLDER'], filename)
            template.save(template_path)
            session['template_path'] = template_path
            session['template_uploaded'] = True
            flash('Template uploaded successfully!', 'success')
        else:
            flash('Please upload a valid .docx file.', 'error')
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('SELECT id, name, usn FROM students ORDER BY name')
    students = c.fetchall()
    conn.close()
    
    return render_template('template_upload.html', students=students)

@app.route('/fill-template', methods=['POST'])
def fill_template():
    selected_ids = request.form.getlist('selected_students')
    template_path = session.get('template_path')
    
    if not selected_ids or not template_path:
        flash('Please select students and upload a template first.', 'error')
        return redirect(url_for('template_upload'))
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    placeholders = ','.join('?' for _ in selected_ids)
    query = f'SELECT * FROM students WHERE id IN ({placeholders})'
    c.execute(query, selected_ids)
    students = c.fetchall()
    conn.close()
    
    doc = Document(template_path)
    
    # Create replacement dictionary for first student
    if students:
        student = students[0]  # Use first selected student
        replacements = {
            '[NAME]': student[1] or '',
            '[USN]': student[6] or '',
            '[DOB]': student[2] or '',
            '[FATHER]': student[4] or '',
            '[MOTHER]': student[3] or '',
            '[BRANCH]': student[5] or '',
            '[PHONE]': student[7] or '',
            '[EMAIL]': student[8] or '',
            '[SPORTS]': student[10] or '',
            '[BLOOD_GROUP]': (student[11] if len(student) > 11 and student[11] else ''),
            '[GENDER]': (student[12] if len(student) > 12 and student[12] else '')
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    filename = f'filled_template_{students[0][1] if students else "document"}.docx'
    return send_file(temp_file.name, as_attachment=True, download_name=filename)

if __name__ == '__main__':
    init_db()
    app.run(debug=True)
