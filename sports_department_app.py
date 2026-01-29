import sys
import os
import sqlite3
from PyQt5 import QtWidgets, QtGui, QtCore
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QLineEdit, QTextEdit, QComboBox, QFileDialog,
    QMessageBox, QTableWidget, QTableWidgetItem, QHeaderView, QCheckBox
)
from PyQt5.QtGui import QPixmap
from PyQt5.QtCore import Qt
from docx import Document

DB_NAME = "students.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            dob TEXT,
            mother_name TEXT,
            father_name TEXT,
            branch_sem TEXT,
            usn TEXT UNIQUE,
            phone TEXT,
            email TEXT,
            photo_path TEXT,
            sports TEXT
        )
    ''')
    conn.commit()
    conn.close()

class DataEntryPage(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.photo_path = None
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        title = QLabel("Student Data Entry")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 24px; font-weight: bold; margin-bottom: 20px;")
        layout.addWidget(title)

        self.name_input = QLineEdit()
        self.dob_input = QLineEdit()
        self.mother_name_input = QLineEdit()
        self.father_name_input = QLineEdit()
        self.branch_sem_input = QLineEdit()
        self.usn_input = QLineEdit()
        self.phone_input = QLineEdit()
        self.email_input = QLineEdit()
        self.sports_input = QLineEdit()

        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        def add_labeled_widget(label_text, widget):
            label = QLabel(label_text)
            label.setStyleSheet("font-size: 14px; font-weight: 600;")
            form_layout.addWidget(label)
            form_layout.addWidget(widget)

        add_labeled_widget("Name:", self.name_input)
        add_labeled_widget("Date of Birth (YYYY-MM-DD):", self.dob_input)
        add_labeled_widget("Mother's Name:", self.mother_name_input)
        add_labeled_widget("Father's Name:", self.father_name_input)
        add_labeled_widget("Branch & Semester:", self.branch_sem_input)
        add_labeled_widget("USN (10 digits):", self.usn_input)
        add_labeled_widget("Phone Number (10 digits):", self.phone_input)
        add_labeled_widget("Email:", self.email_input)
        add_labeled_widget("Sports (comma separated):", self.sports_input)

        photo_layout = QVBoxLayout()
        self.photo_label = QLabel("No photo selected")
        self.photo_label.setAlignment(Qt.AlignCenter)
        self.photo_label.setFixedSize(200, 200)
        self.photo_label.setStyleSheet("border: 1px solid #ccc; background-color: #f9f9f9;")
        photo_btn = QPushButton("Select Photo")
        photo_btn.clicked.connect(self.select_photo)
        photo_layout.addWidget(self.photo_label)
        photo_layout.addWidget(photo_btn)
        photo_layout.setAlignment(Qt.AlignCenter)

        form_layout.addLayout(photo_layout)

        save_btn = QPushButton("Save Student")
        save_btn.setStyleSheet("font-size: 16px; padding: 10px;")
        save_btn.clicked.connect(self.save_student)

        layout.addLayout(form_layout)
        layout.addWidget(save_btn)
        self.setLayout(layout)

    def select_photo(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Select Student Photo", "", "Image Files (*.png *.jpg *.jpeg)", options=options)
        if file_name:
            self.photo_path = file_name
            pixmap = QPixmap(file_name)
            pixmap = pixmap.scaled(self.photo_label.width(), self.photo_label.height(), Qt.KeepAspectRatio)
            self.photo_label.setPixmap(pixmap)

    def save_student(self):
        name = self.name_input.text().strip()
        dob = self.dob_input.text().strip()
        mother_name = self.mother_name_input.text().strip()
        father_name = self.father_name_input.text().strip()
        branch_sem = self.branch_sem_input.text().strip()
        usn = self.usn_input.text().strip()
        phone = self.phone_input.text().strip()
        email = self.email_input.text().strip()
        sports = self.sports_input.text().strip()

        # Validation
        if not name:
            QMessageBox.warning(self, "Validation Error", "Name is required.")
            return
        if len(usn) != 10:
            QMessageBox.warning(self, "Validation Error", "USN must be exactly 10 characters.")
            return
        if not (phone.isdigit() and len(phone) == 10):
            QMessageBox.warning(self, "Validation Error", "Phone number must be numeric and exactly 10 digits.")
            return
        # Additional validations can be added here

        # Save photo to local folder
        photo_dest = None
        if self.photo_path:
            photos_dir = "photos"
            if not os.path.exists(photos_dir):
                os.makedirs(photos_dir)
            photo_filename = os.path.basename(self.photo_path)
            photo_dest = os.path.join(photos_dir, photo_filename)
            if not os.path.exists(photo_dest):
                import shutil
                shutil.copy(self.photo_path, photo_dest)

        # Save to DB
        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                INSERT INTO students (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_dest, sports))
            conn.commit()
            conn.close()
            QMessageBox.information(self, "Success", "Student saved successfully.")
            self.clear_form()
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "Error", "USN must be unique. A student with this USN already exists.")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to save student: {str(e)}")

    def clear_form(self):
        self.name_input.clear()
        self.dob_input.clear()
        self.mother_name_input.clear()
        self.father_name_input.clear()
        self.branch_sem_input.clear()
        self.usn_input.clear()
        self.phone_input.clear()
        self.email_input.clear()
        self.sports_input.clear()
        self.photo_label.clear()
        self.photo_label.setText("No photo selected")
        self.photo_path = None

class EditDataPage(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_student_id = None
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["ID", "Name", "USN", "Phone", "Edit"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)

        layout.addWidget(self.table)
        self.setLayout(layout)
        self.load_data()

    def load_data(self):
        self.table.setRowCount(0)
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("SELECT id, name, usn, phone FROM students")
        rows = c.fetchall()
        conn.close()

        for row_num, row_data in enumerate(rows):
            self.table.insertRow(row_num)
            for col_num, data in enumerate(row_data):
                self.table.setItem(row_num, col_num, QTableWidgetItem(str(data)))
            edit_btn = QPushButton("Edit")
            edit_btn.clicked.connect(lambda checked, r=row_data[0]: self.edit_student(r))
            self.table.setCellWidget(row_num, 4, edit_btn)

    def edit_student(self, student_id):
        self.selected_student_id = student_id
        self.edit_window = EditStudentWindow(student_id)
        self.edit_window.show()
        self.edit_window.saved.connect(self.load_data)

class EditStudentWindow(QWidget):
    saved = QtCore.pyqtSignal()

    def __init__(self, student_id):
        super().__init__()
        self.student_id = student_id
        self.photo_path = None
        self.init_ui()
        self.load_student_data()

    def init_ui(self):
        self.setWindowTitle("Edit Student")
        layout = QVBoxLayout()

        self.name_input = QLineEdit()
        self.dob_input = QLineEdit()
        self.mother_name_input = QLineEdit()
        self.father_name_input = QLineEdit()
        self.branch_sem_input = QLineEdit()
        self.usn_input = QLineEdit()
        self.phone_input = QLineEdit()
        self.email_input = QLineEdit()
        self.sports_input = QLineEdit()

        self.photo_label = QLabel("No photo selected")
        self.photo_label.setAlignment(Qt.AlignCenter)
        self.photo_label.setFixedSize(200, 200)

        photo_btn = QPushButton("Select Photo")
        photo_btn.clicked.connect(self.select_photo)

        form_layout = QVBoxLayout()
        form_layout.addWidget(QLabel("Name:"))
        form_layout.addWidget(self.name_input)
        form_layout.addWidget(QLabel("Date of Birth (YYYY-MM-DD):"))
        form_layout.addWidget(self.dob_input)
        form_layout.addWidget(QLabel("Mother's Name:"))
        form_layout.addWidget(self.mother_name_input)
        form_layout.addWidget(QLabel("Father's Name:"))
        form_layout.addWidget(self.father_name_input)
        form_layout.addWidget(QLabel("Branch & Semester:"))
        form_layout.addWidget(self.branch_sem_input)
        form_layout.addWidget(QLabel("USN (10 digits):"))
        form_layout.addWidget(self.usn_input)
        form_layout.addWidget(QLabel("Phone Number (10 digits):"))
        form_layout.addWidget(self.phone_input)
        form_layout.addWidget(QLabel("Email:"))
        form_layout.addWidget(self.email_input)
        form_layout.addWidget(QLabel("Sports (comma separated):"))
        form_layout.addWidget(self.sports_input)
        form_layout.addWidget(self.photo_label)
        form_layout.addWidget(photo_btn)

        save_btn = QPushButton("Save Changes")
        save_btn.clicked.connect(self.save_changes)

        layout.addLayout(form_layout)
        layout.addWidget(save_btn)
        self.setLayout(layout)

    def load_student_data(self):
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("SELECT name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports FROM students WHERE id=?", (self.student_id,))
        row = c.fetchone()
        conn.close()
        if row:
            self.name_input.setText(row[0])
            self.dob_input.setText(row[1])
            self.mother_name_input.setText(row[2])
            self.father_name_input.setText(row[3])
            self.branch_sem_input.setText(row[4])
            self.usn_input.setText(row[5])
            self.phone_input.setText(row[6])
            self.email_input.setText(row[7])
            self.sports_input.setText(row[9])
            self.photo_path = row[8]
            if self.photo_path and os.path.exists(self.photo_path):
                pixmap = QPixmap(self.photo_path)
                pixmap = pixmap.scaled(self.photo_label.width(), self.photo_label.height(), Qt.KeepAspectRatio)
                self.photo_label.setPixmap(pixmap)
            else:
                self.photo_label.setText("No photo selected")

    def select_photo(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Select Student Photo", "", "Image Files (*.png *.jpg *.jpeg)", options=options)
        if file_name:
            self.photo_path = file_name
            pixmap = QPixmap(file_name)
            pixmap = pixmap.scaled(self.photo_label.width(), self.photo_label.height(), Qt.KeepAspectRatio)
            self.photo_label.setPixmap(pixmap)

    def save_changes(self):
        name = self.name_input.text().strip()
        dob = self.dob_input.text().strip()
        mother_name = self.mother_name_input.text().strip()
        father_name = self.father_name_input.text().strip()
        branch_sem = self.branch_sem_input.text().strip()
        usn = self.usn_input.text().strip()
        phone = self.phone_input.text().strip()
        email = self.email_input.text().strip()
        sports = self.sports_input.text().strip()

        if not name:
            QMessageBox.warning(self, "Validation Error", "Name is required.")
            return
        if len(usn) != 10:
            QMessageBox.warning(self, "Validation Error", "USN must be exactly 10 characters.")
            return
        if not (phone.isdigit() and len(phone) == 10):
            QMessageBox.warning(self, "Validation Error", "Phone number must be numeric and exactly 10 digits.")
            return

        photo_dest = None
        if self.photo_path:
            photos_dir = "photos"
            if not os.path.exists(photos_dir):
                os.makedirs(photos_dir)
            photo_filename = os.path.basename(self.photo_path)
            photo_dest = os.path.join(photos_dir, photo_filename)
            if not os.path.exists(photo_dest) or photo_dest != self.photo_path:
                import shutil
                shutil.copy(self.photo_path, photo_dest)

        try:
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute('''
                UPDATE students SET name=?, dob=?, mother_name=?, father_name=?, branch_sem=?, usn=?, phone=?, email=?, photo_path=?, sports=?
                WHERE id=?
            ''', (name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_dest, sports, self.student_id))
            conn.commit()
            conn.close()
            QMessageBox.information(self, "Success", "Student updated successfully.")
            self.saved.emit()
            self.close()
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "Error", "USN must be unique. A student with this USN already exists.")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to update student: {str(e)}")

class SelectStudentsPage(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_students = set()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["Select", "ID", "Name", "USN", "Phone", "Email"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)

        layout.addWidget(self.table)

        self.setLayout(layout)
        self.load_data()

    def load_data(self):
        self.table.setRowCount(0)
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("SELECT id, name, usn, phone, email FROM students")
        rows = c.fetchall()
        conn.close()

        for row_num, row_data in enumerate(rows):
            self.table.insertRow(row_num)

            checkbox = QCheckBox()
            checkbox.stateChanged.connect(lambda state, r=row_data[0]: self.on_checkbox_state_changed(state, r))
            self.table.setCellWidget(row_num, 0, checkbox)

            for col_num, data in enumerate(row_data):
                self.table.setItem(row_num, col_num + 1, QTableWidgetItem(str(data)))

    def on_checkbox_state_changed(self, state, student_id):
        if state == Qt.Checked:
            self.selected_students.add(student_id)
        else:
            self.selected_students.discard(student_id)

class GenerateReportPage(QWidget):
    def __init__(self, select_students_page, parent=None):
        super().__init__(parent)
        self.select_students_page = select_students_page
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.info_label = QLabel("Select students from the 'Select Students' page, then click 'Generate Report' to create a Word document.")
        layout.addWidget(self.info_label)

        self.generate_btn = QPushButton("Generate Report")
        self.generate_btn.clicked.connect(self.generate_report)
        layout.addWidget(self.generate_btn)

        self.setLayout(layout)

    def generate_report(self):
        selected_ids = self.select_students_page.selected_students
        if not selected_ids:
            QMessageBox.warning(self, "No Selection", "Please select at least one student to generate the report.")
            return

        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        placeholders = ",".join("?" for _ in selected_ids)
        query = f"SELECT name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports FROM students WHERE id IN ({placeholders})"
        c.execute(query, tuple(selected_ids))
        students = c.fetchall()
        conn.close()

        doc = Document()
        doc.add_heading("Selected Students Report", 0)

        for student in students:
            name, dob, mother_name, father_name, branch_sem, usn, phone, email, photo_path, sports = student
            doc.add_heading(name, level=1)
            doc.add_paragraph(f"Date of Birth: {dob}")
            doc.add_paragraph(f"Mother's Name: {mother_name}")
            doc.add_paragraph(f"Father's Name: {father_name}")
            doc.add_paragraph(f"Branch & Semester: {branch_sem}")
            doc.add_paragraph(f"USN: {usn}")
            doc.add_paragraph(f"Phone: {phone}")
            doc.add_paragraph(f"Email: {email}")
            doc.add_paragraph(f"Sports: {sports}")

            if photo_path and os.path.exists(photo_path):
                try:
                    doc.add_picture(photo_path, width=QtCore.QSize(150, 150))
                except Exception as e:
                    doc.add_paragraph(f"[Error loading photo: {str(e)}]")

            doc.add_page_break()

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Report", "", "Word Documents (*.docx)")
        if save_path:
            doc.save(save_path)
            QMessageBox.information(self, "Report Generated", f"Report saved successfully at:\n{save_path}")

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("College Sports Department")
        self.setGeometry(100, 100, 800, 600)

        self.data_entry_page = DataEntryPage()
        self.edit_data_page = EditDataPage()
        self.select_students_page = SelectStudentsPage()
        self.generate_report_page = GenerateReportPage(self.select_students_page)

        self.init_ui()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        layout = QVBoxLayout()

        btn_data_entry = QPushButton("Data Entry")
        btn_data_entry.clicked.connect(self.show_data_entry)

        btn_edit_data = QPushButton("Edit Data")
        btn_edit_data.clicked.connect(self.show_edit_data)

        btn_select_students = QPushButton("Select Students")
        btn_select_students.clicked.connect(self.show_select_students)

        btn_generate_report = QPushButton("Generate Report")
        btn_generate_report.clicked.connect(self.show_generate_report)

        btn_layout = QHBoxLayout()
        btn_layout.addWidget(btn_data_entry)
        btn_layout.addWidget(btn_edit_data)
        btn_layout.addWidget(btn_select_students)
        btn_layout.addWidget(btn_generate_report)

        layout.addLayout(btn_layout)
        layout.addWidget(self.data_entry_page)
        layout.addWidget(self.edit_data_page)
        layout.addWidget(self.select_students_page)
        layout.addWidget(self.generate_report_page)

        self.edit_data_page.hide()
        self.select_students_page.hide()
        self.generate_report_page.hide()

        central_widget.setLayout(layout)

    def show_data_entry(self):
        self.data_entry_page.show()
        self.edit_data_page.hide()
        self.select_students_page.hide()
        self.generate_report_page.hide()

    def show_edit_data(self):
        self.edit_data_page.load_data()
        self.edit_data_page.show()
        self.data_entry_page.hide()
        self.select_students_page.hide()
        self.generate_report_page.hide()

    def show_select_students(self):
        self.select_students_page.load_data()
        self.select_students_page.show()
        self.data_entry_page.hide()
        self.edit_data_page.hide()
        self.generate_report_page.hide()

    def show_generate_report(self):
        self.generate_report_page.show()
        self.data_entry_page.hide()
        self.edit_data_page.hide()
        self.select_students_page.hide()

def main():
    init_db()
    app = QApplication(sys.argv)
    main_win = MainWindow()
    main_win.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
