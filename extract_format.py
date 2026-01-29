from docx import Document

# Read the VTU ELE.docx file
doc = Document('REPORT FORMATS/VTU ELE.docx')

print("=== VTU ELE.docx Content ===")
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip():
        print(f"Para {i}: '{paragraph.text}'")
        print(f"Alignment: {paragraph.alignment}")
        for run in paragraph.runs:
            print(f"  Run: '{run.text}' - Bold: {run.bold}, Underline: {run.underline}")
        print()

# Check tables
if doc.tables:
    print("\n=== Tables ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print(f"  Cell: '{cell.text}'")